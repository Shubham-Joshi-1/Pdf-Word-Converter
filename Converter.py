import os
import subprocess
from docx import Document
from docx.shared import Pt
from pdf2docx import Converter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import pdfplumber
import fitz  # PyMuPDF


def convert_docx_to_pdf(input_docx, output_pdf):
    """Convert a Word document (.docx) to a PDF."""
    from docx2pdf import convert
    try:
        convert(input_docx, output_pdf)
        print(f"Converted {input_docx} to {output_pdf}.")
    except Exception as e:
        print(f"Error converting {input_docx} to PDF: {e}")


def convert_pdf_to_docx(input_pdf, output_docx):
    """Convert a PDF to a Word document (.docx)."""
    try:
        converter = Converter(input_pdf)
        converter.convert(output_docx, start=0, end=None)
        converter.close()
        print(f"Converted {input_pdf} to {output_docx}.")
    except Exception as e:
        print(f"Error converting {input_pdf} to DOCX: {e}")


def extract_fonts_from_pdf(input_pdf):
    """Extract font information from a PDF."""
    try:
        doc = fitz.open(input_pdf)
        font_details = {}
        for page_num, page in enumerate(doc):
            fonts = page.get_fonts(full=True)
            font_details[page_num] = fonts
        print(f"Extracted fonts from {input_pdf}: {font_details}")
    except Exception as e:
        print(f"Error extracting fonts from {input_pdf}: {e}")


def extract_tables_from_pdf(input_pdf):
    """Extract tables from a PDF and print them."""
    try:
        with pdfplumber.open(input_pdf) as pdf:
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                for table in tables:
                    print(f"Table from page {i + 1}: {table}")
    except Exception as e:
        print(f"Error extracting tables from {input_pdf}: {e}")


def create_word_with_styles(output_docx, text, font_name="Arial", font_size=12):
    """Create a Word document with specified styles."""
    try:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        document.save(output_docx)
        print(f"Created Word document {output_docx} with styles.")
    except Exception as e:
        print(f"Error creating Word document {output_docx}: {e}")


def split_pdf(input_pdf, output_dir):
    """Split a PDF into individual pages."""
    from PyPDF2 import PdfReader, PdfWriter
    try:
        reader = PdfReader(input_pdf)
        for i, page in enumerate(reader.pages):
            writer = PdfWriter()
            writer.add_page(page)
            output_path = os.path.join(output_dir, f"page_{i + 1}.pdf")
            with open(output_path, "wb") as output_file:
                writer.write(output_file)
            print(f"Extracted page {i + 1} to {output_path}.")
    except Exception as e:
        print(f"Error splitting {input_pdf}: {e}")


def main():
    print("Select an option:")
    print("1. Convert DOCX to PDF")
    print("2. Convert PDF to DOCX")
    print("3. Extract Fonts from PDF")
    print("4. Extract Tables from PDF")
    print("5. Split PDF into Pages")
    
    choice = input("Enter your choice (1-5): ")
    
    if choice == "1":
        input_docx = input("Enter the path to the DOCX file: ")
        output_pdf = input("Enter the path to save the PDF file: ")
        convert_docx_to_pdf(input_docx, output_pdf)
    elif choice == "2":
        input_pdf = input("Enter the path to the PDF file: ")
        output_docx = input("Enter the path to save the DOCX file: ")
        convert_pdf_to_docx(input_pdf, output_docx)
    elif choice == "3":
        input_pdf = input("Enter the path to the PDF file: ")
        extract_fonts_from_pdf(input_pdf)
    elif choice == "4":
        input_pdf = input("Enter the path to the PDF file: ")
        extract_tables_from_pdf(input_pdf)
    elif choice == "5":
        input_pdf = input("Enter the path to the PDF file: ")
        output_dir = input("Enter the directory to save the split pages: ")
        split_pdf(input_pdf, output_dir)
    else:
        print("Invalid choice. Exiting.")


if __name__ == "__main__":
    main()
